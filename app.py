from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from urllib.parse import urlparse
import cgi
import csv
import email
import hashlib
import io
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
import time
import zipfile
from datetime import datetime
from html.parser import HTMLParser

from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from PIL import Image, ImageEnhance, ImageFilter, ImageOps
from pypdf import PdfReader
from docx import Document


ROOT = Path(__file__).resolve().parent
PUBLIC = ROOT / "public"
OUTPUTS = ROOT / "outputs"
PENDING = ROOT / "pending_reviews"
PROCESSED = ROOT / "processed_fax"
REJECTED = ROOT / "rejected_fax"
PC_EXCEL_DIR = ROOT / "pc_excel"
CONFIG_PATH = ROOT / "config.json"
HISTORY_PATH = ROOT / "history.jsonl"
APP_VERSION = "20sec-accuracy-20260531-01"

DEFAULT_CONFIG = {
    "incoming_fax_dir": str(ROOT / "incoming_fax"),
    "ledger_path": str(PC_EXCEL_DIR / "fax_ledger.xlsx"),
    "reviewer_default": "社員確認",
    "auto_process_to_review": True,
    "ocr_min_chars": 8,
    "ocr_psm": "6",
    "ocr_timeout_seconds": 8,
    "ocr_preprocess": True,
    "ocr_preprocess_mode": "clean",
    "ocr_processed_format": "jpeg",
    "ocr_jpeg_quality": 82,
    "ocr_pdf_dpi": 140,
    "ocr_pdf_image_format": "jpeg",
    "ocr_target_long_side": 1100,
    "ocr_min_long_side": 800,
    "ocr_max_pages": 1,
    "ocr_retry_when_empty": True,
    "request_timeout_seconds": 19,
    "max_upload_files": 1,
    "pdf_text_max_pages": 2,
    "xlsx_max_rows_per_sheet": 100,
    "docx_max_paragraphs": 70,
    "docx_max_table_rows": 70,
    "html_max_chars": 24000,
    "email_max_parts": 4,
    "zip_max_files": 5,
    "zip_max_file_bytes": 1500000,
    "text_max_chars": 24000,
}


FIELD_LABELS = {
    "received_date": "日付",
    "document_type": "書類種別",
    "sender_company": "送信元会社",
    "recipient_company": "宛先会社",
    "fax_no": "FAX番号",
    "order_no": "注文番号",
    "project_name": "工事名",
    "site_name": "現場名",
    "delivery_place": "納入場所",
    "delivery_date": "納期",
    "desired_time": "希望時間",
    "person_in_charge": "担当者",
    "notes": "備考",
}

ITEM_LABELS = {
    "item_code": "品番",
    "item_name": "商品名・資材名",
    "quantity": "数量",
    "unit": "単位",
    "unit_price": "単価",
    "amount": "金額",
}

LEDGER_COLUMNS = {
    "review_status": "確認状況",
    "reviewer": "確認者",
    "reviewed_at": "確認日時",
    "processed_at": "入力日時",
    "source_name": "元ファイル",
    **FIELD_LABELS,
    **ITEM_LABELS,
}

IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".bmp", ".tif", ".tiff", ".webp", ".heic", ".heif"}
TEXT_EXTENSIONS = {".txt", ".csv", ".tsv", ".log", ".json", ".xml", ".md", ".rtf"}
PDF_EXTENSIONS = {".pdf"}
OFFICE_EXTENSIONS = {".docx", ".xlsx", ".xlsm"}
HTML_EXTENSIONS = {".html", ".htm"}
EMAIL_EXTENSIONS = {".eml"}
ARCHIVE_EXTENSIONS = {".zip"}


def now_text():
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


def load_config():
    config = DEFAULT_CONFIG.copy()
    if CONFIG_PATH.exists():
        try:
            config.update(json.loads(CONFIG_PATH.read_text(encoding="utf-8")))
        except json.JSONDecodeError:
            pass
    if os.environ.get("INCOMING_FAX_DIR"):
        config["incoming_fax_dir"] = os.environ["INCOMING_FAX_DIR"]
    if os.environ.get("LEDGER_PATH"):
        config["ledger_path"] = os.environ["LEDGER_PATH"]
    return config


def save_config(config):
    CONFIG_PATH.write_text(json.dumps(config, ensure_ascii=False, indent=2), encoding="utf-8")


def incoming_dir():
    return Path(load_config()["incoming_fax_dir"])


def ledger_path():
    return Path(load_config()["ledger_path"])


def write_history(event, data):
    entry = {"time": now_text(), "event": event, **data}
    HISTORY_PATH.parent.mkdir(exist_ok=True)
    with HISTORY_PATH.open("a", encoding="utf-8") as handle:
        handle.write(json.dumps(entry, ensure_ascii=False) + "\n")


def normalize_text(text):
    replacements = {
        "\r\n": "\n",
        "\r": "\n",
        "　": " ",
        "㈱": "株式会社",
        "（株）": "株式会社",
        "／": "/",
        "－": "-",
        "ー": "-",
        "―": "-",
        "：": ":",
        "，": ",",
    }
    for before, after in replacements.items():
        text = text.replace(before, after)
    text = text.translate(str.maketrans("０１２３４５６７８９ＡＢＣＤＥＦＧＨＩＪＫＬＭＮＯＰＱＲＳＴＵＶＷＸＹＺａｂｃｄｅｆｇｈｉｊｋｌｍｎｏｐｑｒｓｔｕｖｗｘｙｚ", "0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz"))
    return re.sub(r"[ \t]+", " ", text).strip()


def first_match(patterns, text, default=""):
    for pattern in patterns:
        match = re.search(pattern, text, re.MULTILINE | re.IGNORECASE)
        if match:
            return match.group(1).strip(" :：\t")
    return default


def normalize_date(value):
    if not value:
        return ""
    value = value.strip()
    match = re.search(r"(\d{4})[./年-](\d{1,2})[./月-](\d{1,2})", value)
    if match:
        y, m, d = match.groups()
        return f"{int(y):04d}-{int(m):02d}-{int(d):02d}"
    match = re.search(r"(\d{1,2})[./月-](\d{1,2})", value)
    if match:
        m, d = match.groups()
        return f"{datetime.now().year:04d}-{int(m):02d}-{int(d):02d}"
    return value


def number_or_blank(value):
    if value is None:
        return ""
    cleaned = re.sub(r"[^0-9.-]", "", str(value))
    if cleaned in {"", "-", ".", "-."}:
        return ""
    try:
        number = float(cleaned)
        return int(number) if number.is_integer() else number
    except ValueError:
        return ""


def parse_items(text):
    items = []
    line_patterns = [
        re.compile(
            r"(?P<code>[A-Z0-9][A-Z0-9\-_./]{1,})\s+"
            r"(?P<name>.+?)\s+"
            r"(?P<qty>\d+(?:\.\d+)?)\s*(?P<unit>個|枚|箱|本|台|kg|t|式|袋|缶|束|セット|m|ｍ|m2|m3|㎡|㎥)?\s+"
            r"(?P<price>[\d,]+)?\s+"
            r"(?P<amount>[\d,]+)?$",
            re.IGNORECASE,
        ),
        re.compile(
            r"(?:品番|型番|コード)[: ]?(?P<code>[A-Z0-9\-_.\/]+).*?"
            r"(?:商品名|品名|資材名|名称)[: ]?(?P<name>.+?)\s+"
            r"(?:数量|数|個数)[: ]?(?P<qty>\d+(?:\.\d+)?).*?"
            r"(?:単価[: ]?(?P<price>[\d,]+))?.*?"
            r"(?:(?:金額|小計)[: ]?(?P<amount>[\d,]+))?",
            re.IGNORECASE,
        ),
        re.compile(
            r"(?P<name>.+?)\s+"
            r"(?P<qty>\d+(?:\.\d+)?)\s*(?P<unit>個|枚|箱|本|台|kg|t|式|袋|缶|束|セット|m|ｍ|m2|m3|㎡|㎥)\s+"
            r"(?P<price>[\d,]+)?\s+"
            r"(?P<amount>[\d,]+)?$",
            re.IGNORECASE,
        ),
    ]

    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if any(word in line for word in ["品番", "商品名", "数量", "単価", "金額"]) and len(line) < 18:
            continue
        for pattern in line_patterns:
            match = pattern.search(line)
            if not match:
                continue
            data = match.groupdict()
            qty = number_or_blank(data.get("qty"))
            price = number_or_blank(data.get("price"))
            amount = number_or_blank(data.get("amount"))
            if amount == "" and qty != "" and price != "":
                amount = qty * price
            items.append(
                {
                    "item_code": data.get("code", "").strip(),
                    "item_name": data.get("name", "").strip(" -:："),
                    "quantity": qty,
                    "unit": (data.get("unit") or "個").strip(),
                    "unit_price": price,
                    "amount": amount,
                }
            )
            break

    if not items:
        items.append({"item_code": "", "item_name": "", "quantity": "", "unit": "", "unit_price": "", "amount": ""})
    return items


def extract_fax(text):
    text = normalize_text(text)
    header = {
        "received_date": normalize_date(first_match([r"(?:受信日|受付日|日付|発注日|作成日)[: ]*([0-9./年月日-]+)"], text)),
        "document_type": first_match(
            [r"(見積依頼書|発注書|注文書|納品依頼書|資材依頼書|作業依頼書)", r"(?:書類種別|帳票種別)[:： ]*(.+)"],
            text,
        ),
        "sender_company": first_match(
            [
                r"(?:送信元|差出人|From|発注元|依頼元|会社名)[: ]*(.+)",
                r"(.+?(?:株式会社|有限会社|合同会社|商事|工業|産業|物流|建設|工務店|土木|設備))\s*(?:御中|様)?$",
            ],
            text,
        ),
        "recipient_company": first_match([r"(?:宛先|To|納入先|提出先)[: ]*(.+)"], text),
        "fax_no": first_match([r"(?:FAX|Fax|ファックス)[: ]*([0-9\-\(\) ]{8,})"], text),
        "order_no": first_match([r"(?:注文番号|発注番号|注文No\.?|Order No\.?|No\.?)[: ]*([A-Z0-9\-_.\/]+)"], text),
        "project_name": first_match([r"(?:工事名|案件名|件名|物件名|プロジェクト名)[: ]*(.+)"], text),
        "site_name": first_match([r"(?:現場名|現場|作業所名|作業所)[: ]*(.+)"], text),
        "delivery_place": first_match([r"(?:納入場所|搬入場所|納品場所|現場住所|住所|配送先|届け先)[: ]*(.+)"], text),
        "delivery_date": normalize_date(first_match([r"(?:納期|希望納期|納入日|搬入日|納品日|配達日|希望日)[: ]*([0-9./年月日-]+)"], text)),
        "desired_time": first_match([r"(?:希望時間|搬入時間|納品時間|時間指定|時間帯)[: ]*(.+)"], text),
        "person_in_charge": first_match([r"(?:担当者|担当|現場担当|連絡先担当)[: ]*(.+)"], text),
        "notes": first_match([r"(?:備考|摘要|連絡事項|注意事項|コメント)[: ]*(.+)"], text),
    }
    confidence = {key: 0.92 if value else 0.35 for key, value in header.items()}
    return {"header": header, "items": parse_items(text), "confidence": confidence, "raw_text": text}


def read_pdf_text(file_bytes):
    reader = PdfReader(io.BytesIO(file_bytes))
    pages = []
    max_pages = int(load_config().get("pdf_text_max_pages", 3))
    for index, page in enumerate(reader.pages, start=1):
        if index > max_pages:
            break
        text = page.extract_text() or ""
        if text.strip():
            pages.append(f"--- PDF {index}ページ ---\n{text.strip()}")
    text = "\n\n".join(pages).strip()
    return text or read_scanned_pdf_text(file_bytes)


def read_scanned_pdf_text(file_bytes):
    pdftoppm = shutil.which("pdftoppm")
    tesseract = shutil.which("tesseract")
    if not pdftoppm or not tesseract:
        return ""

    with tempfile.TemporaryDirectory() as temp_dir:
        temp = Path(temp_dir)
        pdf_path = temp / "input.pdf"
        image_base = temp / "page"
        pdf_path.write_bytes(file_bytes)
        config = load_config()
        dpi = str(config.get("ocr_pdf_dpi", 180))
        max_pages = str(config.get("ocr_max_pages", 2))
        image_format = str(config.get("ocr_pdf_image_format", "jpeg")).lower()
        format_args = ["-jpeg", "-jpegopt", "quality=85"] if image_format in {"jpg", "jpeg"} else ["-png"]
        convert = subprocess.run(
            [pdftoppm, *format_args, "-r", dpi, "-f", "1", "-l", max_pages, str(pdf_path), str(image_base)],
            capture_output=True,
            text=True,
            timeout=max(3, int(config.get("ocr_timeout_seconds", 2)) + 2),
        )
        if convert.returncode != 0:
            return ""

        pages = []
        page_images = sorted(path for path in temp.glob("page-*.*") if path.suffix.lower() in {".jpg", ".jpeg", ".png", ".ppm"})
        for index, image_path in enumerate(page_images, start=1):
            ocr_image = preprocess_image_file(image_path) if load_config().get("ocr_preprocess", True) else image_path
            page_text, _ = run_fast_ocr(tesseract, ocr_image, temp / f"ocr_{index}")
            if page_text:
                pages.append(f"--- OCR PDF {index}ページ ---\n{page_text}")
        return "\n\n".join(pages).strip()


def meaningful_text_length(text):
    return len(re.sub(r"\s+", "", text or ""))


def japanese_char_count(text):
    return len(re.findall(r"[\u3040-\u30ff\u3400-\u9fff]", text or ""))


def english_signal_count(text):
    return len(re.findall(r"[A-Za-z0-9]", text or ""))


def run_tesseract(tesseract, image_path, output_base, lang):
    config = load_config()
    command = [
        tesseract,
        str(image_path),
        str(output_base),
        "-l",
        lang,
        "--oem",
        "1",
        "--psm",
        str(config.get("ocr_psm", "6")),
        "-c",
        "load_system_dawg=0",
        "-c",
        "load_freq_dawg=0",
        "-c",
        f"debug_file={'NUL' if os.name == 'nt' else '/dev/null'}",
        "-c",
        "tessedit_do_invert=0",
    ]
    started = datetime.now()
    result = subprocess.run(
        command,
        capture_output=True,
        text=True,
        timeout=int(config.get("ocr_timeout_seconds", 45)),
    )
    elapsed_ms = int((datetime.now() - started).total_seconds() * 1000)
    if result.returncode != 0:
        return "", f"{lang}: {result.stderr.strip() or result.stdout.strip()}", elapsed_ms
    text_path = output_base.with_suffix(".txt")
    text = text_path.read_text(encoding="utf-8", errors="ignore") if text_path.exists() else ""
    return text.strip(), "", elapsed_ms


def run_fast_ocr(tesseract, image_path, output_base):
    config = load_config()
    min_chars = int(config.get("ocr_min_chars", 12))
    retry_when_empty = bool(config.get("ocr_retry_when_empty", True))

    # Keep production scans fast: one Japanese OCR pass first. Japanese traineddata also
    # handles many alphanumeric fields, so this avoids doing English + Japanese twice.
    candidates = []
    text, warning, elapsed_ms = run_tesseract(tesseract, image_path, output_base, "jpn")
    write_history("ocr_attempt", {"lang": "jpn", "elapsed_ms": elapsed_ms, "chars": meaningful_text_length(text)})
    candidates.append((text, warning, "jpn"))
    if meaningful_text_length(text) >= min_chars:
        return text, warning
    if meaningful_text_length(text) > 0 or not retry_when_empty:
        return text, warning

    for lang in ["jpn+eng"]:
        next_text, next_warning, next_elapsed_ms = run_tesseract(
            tesseract,
            image_path,
            output_base.with_name(output_base.name + "_" + re.sub(r"[^A-Za-z0-9]+", "_", lang)),
            lang,
        )
        write_history("ocr_attempt", {"lang": lang, "elapsed_ms": next_elapsed_ms, "chars": meaningful_text_length(next_text)})
        candidates.append((next_text, next_warning, lang))
        if meaningful_text_length(next_text) >= min_chars and japanese_char_count(next_text) > 0:
            break

    best_text, best_warning, _ = max(candidates, key=lambda item: meaningful_text_length(item[0]))
    return best_text, best_warning


def preprocess_image_file(image_path):
    try:
        image = Image.open(image_path)
        image = ImageOps.exif_transpose(image)
        image = image.convert("L")
        width, height = image.size
        config = load_config()
        mode = str(config.get("ocr_preprocess_mode", "resize")).lower()
        target = int(config.get("ocr_target_long_side", 1500))
        min_long_side = int(config.get("ocr_min_long_side", 0))
        long_side = max(width, height)
        should_resize_down = long_side > target
        should_resize_up = min_long_side > 0 and long_side < min_long_side
        if should_resize_down or should_resize_up:
            resize_to = target if should_resize_down else min_long_side
            scale = resize_to / long_side
            image = image.resize((int(width * scale), int(height * scale)), Image.Resampling.BILINEAR)
        if mode == "clean":
            image = ImageOps.autocontrast(image)
            image = ImageEnhance.Contrast(image).enhance(1.6)
            image = image.filter(ImageFilter.SHARPEN)
            image = image.point(lambda pixel: 255 if pixel > 170 else 0)
        processed_format = str(config.get("ocr_processed_format", "jpeg")).lower()
        if processed_format in {"jpg", "jpeg"}:
            processed = image_path.with_name(image_path.stem + "_ocr.jpg")
            image.save(processed, "JPEG", quality=int(config.get("ocr_jpeg_quality", 72)), optimize=False)
        else:
            processed = image_path.with_name(image_path.stem + "_ocr.png")
            image.save(processed)
        return processed
    except Exception as exc:
        write_history("image_preprocess_failed", {"path": str(image_path), "error": str(exc)})
        return image_path


def looks_like_complete_business_text(text):
    if not text:
        return False
    signals = ["FAX", "No", "ORDER", "DATE", "QTY", "TEL", "〒"]
    return meaningful_text_length(text) >= 40 and sum(1 for signal in signals if signal.lower() in text.lower()) >= 2


def read_image_text(file_bytes, suffix):
    if suffix in {".heic", ".heif"}:
        return "", "HEIC/HEIFはこの環境では未対応です。iPhone側で「互換性優先」にするか、JPG/PNGに変換してアップロードしてください。"
    tesseract = shutil.which("tesseract")
    if not tesseract:
        return "", "写真・画像OCRにはTesseract OCRの追加が必要です。RenderのDocker構成では自動で入ります。"

    with tempfile.TemporaryDirectory() as temp_dir:
        temp = Path(temp_dir)
        image_path = temp / f"input{suffix}"
        output_base = temp / "ocr"
        image_path.write_bytes(file_bytes)
        ocr_image = preprocess_image_file(image_path) if load_config().get("ocr_preprocess", True) else image_path
        text, warning = run_fast_ocr(tesseract, ocr_image, output_base)
        if warning and not text:
            return "", f"OCRに失敗しました: {warning}"
        return text.strip(), "" if text.strip() else "画像から文字を読み取れませんでした。写真の明るさや傾きを確認してください。"


class TextOnlyHTMLParser(HTMLParser):
    def __init__(self):
        super().__init__()
        self.parts = []

    def handle_data(self, data):
        text = data.strip()
        if text:
            self.parts.append(text)


def read_docx_text(file_bytes):
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp:
        temp.write(file_bytes)
        temp_path = temp.name
    try:
        config = load_config()
        max_paragraphs = int(config.get("docx_max_paragraphs", 80))
        max_table_rows = int(config.get("docx_max_table_rows", 80))
        document = Document(temp_path)
        parts = []
        for paragraph in document.paragraphs[:max_paragraphs]:
            if paragraph.text.strip():
                parts.append(paragraph.text)
        table_rows = 0
        for table in document.tables:
            for row in table.rows:
                if table_rows >= max_table_rows:
                    return "\n".join(parts).strip()
                values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if values:
                    parts.append(" ".join(values))
                    table_rows += 1
        return "\n".join(parts).strip()
    finally:
        Path(temp_path).unlink(missing_ok=True)


def read_xlsx_text(file_bytes):
    with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as temp:
        temp.write(file_bytes)
        temp_path = temp.name
    workbook = None
    try:
        workbook = load_workbook(temp_path, data_only=True, read_only=True)
        parts = []
        max_rows = int(load_config().get("xlsx_max_rows_per_sheet", 120))
        for sheet in workbook.worksheets:
            parts.append(f"--- Excel sheet: {sheet.title} ---")
            for index, row in enumerate(sheet.iter_rows(values_only=True), start=1):
                if index > max_rows:
                    parts.append(f"--- {sheet.title}: 先頭{max_rows}行まで読み込み ---")
                    break
                values = [str(value) for value in row if value not in (None, "")]
                if values:
                    parts.append(" ".join(values))
        return "\n".join(parts).strip()
    finally:
        if workbook:
            workbook.close()
        Path(temp_path).unlink(missing_ok=True)


def read_html_text(file_bytes):
    max_chars = int(load_config().get("html_max_chars", 30000))
    parser = TextOnlyHTMLParser()
    parser.feed(file_bytes[: max_chars * 4].decode("utf-8", errors="ignore"))
    return "\n".join(parser.parts).strip()[:max_chars]


def read_email_text(file_bytes):
    message = email.message_from_bytes(file_bytes)
    parts = []
    max_parts = int(load_config().get("email_max_parts", 5))
    subject = message.get("subject", "")
    if subject:
        parts.append(f"件名: {subject}")
    processed = 0
    for part in message.walk():
        if processed >= max_parts:
            break
        content_type = part.get_content_type()
        if content_type not in {"text/plain", "text/html"}:
            continue
        payload = part.get_payload(decode=True)
        if not payload:
            continue
        charset = part.get_content_charset() or "utf-8"
        text = payload.decode(charset, errors="ignore")
        if content_type == "text/html":
            text = read_html_text(text.encode("utf-8"))
        if text.strip():
            parts.append(text.strip())
            processed += 1
    return "\n\n".join(parts).strip()


def read_archive_text(file_bytes):
    parts = []
    warnings = []
    config = load_config()
    max_files = int(config.get("zip_max_files", 8))
    max_file_bytes = int(config.get("zip_max_file_bytes", 2000000))
    with zipfile.ZipFile(io.BytesIO(file_bytes)) as archive:
        processed = 0
        for info in archive.infolist():
            if info.is_dir():
                continue
            if processed >= max_files:
                warnings.append(f"ZIP内は先頭{max_files}ファイルまで読み込みました。")
                break
            inner_name = info.filename
            inner_suffix = Path(inner_name).suffix.lower()
            if info.file_size > max_file_bytes:
                warnings.append(f"{inner_name}: skipped because it is larger than {max_file_bytes} bytes")
                continue
            if inner_suffix in ARCHIVE_EXTENSIONS:
                warnings.append(f"{inner_name}: ZIPの中のZIPは読み飛ばしました。")
                continue
            try:
                inner_text, warning = read_file_text(inner_name, archive.read(info))
            except Exception as exc:
                inner_text, warning = "", f"{inner_name}: 読み取りに失敗しました: {exc}"
            if inner_text:
                parts.append(f"--- {inner_name} ---\n{inner_text}")
                processed += 1
            if warning:
                warnings.append(f"{inner_name}: {warning}")
    return "\n\n".join(parts).strip(), " / ".join(warnings)


def read_file_text(filename, file_bytes):
    suffix = Path(filename).suffix.lower()
    if suffix in PDF_EXTENSIONS:
        text = read_pdf_text(file_bytes)
        warning = "" if text else "PDF内に文字情報が見つかりませんでした。スキャンPDFの場合はOCRの追加が必要です。"
        return text, warning
    if suffix in IMAGE_EXTENSIONS:
        return read_image_text(file_bytes, suffix)
    if suffix == ".docx":
        return read_docx_text(file_bytes), ""
    if suffix in {".xlsx", ".xlsm"}:
        return read_xlsx_text(file_bytes), ""
    if suffix == ".doc":
        return "", "古いWord形式（.doc）は未対応です。.docxかPDFに変換してください。"
    if suffix == ".xls":
        return "", "古いExcel形式（.xls）は未対応です。.xlsxかPDFに変換してください。"
    if suffix in HTML_EXTENSIONS:
        return read_html_text(file_bytes), ""
    if suffix in EMAIL_EXTENSIONS:
        return read_email_text(file_bytes), ""
    if suffix in ARCHIVE_EXTENSIONS:
        return read_archive_text(file_bytes)
    if suffix in TEXT_EXTENSIONS:
        max_chars = int(load_config().get("text_max_chars", 30000))
        text = file_bytes[: max_chars * 4].decode("utf-8-sig", errors="ignore").strip()
        warning = "" if len(text) <= max_chars else f"テキストは先頭{max_chars}文字まで読み込みました。"
        return text[:max_chars], warning
    return "", f"{suffix or '拡張子なし'} は未対応です。PDF、画像、Word、Excel、ZIP、テキストに変換してください。"


def create_workbook(payload):
    OUTPUTS.mkdir(exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_source = re.sub(r"[^A-Za-z0-9_-]+", "_", payload.get("source_name", "")).strip("_")
    name = f"fax_excel_{timestamp}{'_' + safe_source[:32] if safe_source else ''}.xlsx"
    output_path = OUTPUTS / name

    wb = Workbook()
    ws = wb.active
    ws.title = "FAX入力"
    raw = wb.create_sheet("原文")

    title_fill = PatternFill("solid", fgColor="1F4E79")
    header_fill = PatternFill("solid", fgColor="D9EAF7")
    thin = Side(style="thin", color="B7C9D6")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    ws["A1"] = "建設FAX入力結果"
    ws["A1"].font = Font(size=16, bold=True, color="FFFFFF")
    ws["A1"].fill = title_fill
    ws.merge_cells("A1:H1")

    row = 3
    header = payload.get("header", {})
    for key, label in FIELD_LABELS.items():
        ws.cell(row=row, column=1, value=label)
        ws.cell(row=row, column=2, value=header.get(key, ""))
        ws.cell(row=row, column=1).fill = header_fill
        row += 1

    row += 1
    start_row = row
    for col, label in enumerate(ITEM_LABELS.values(), start=1):
        cell = ws.cell(row=row, column=col, value=label)
        cell.fill = header_fill
        cell.font = Font(bold=True)
    row += 1

    for item in payload.get("items", []):
        for col, key in enumerate(ITEM_LABELS.keys(), start=1):
            ws.cell(row=row, column=col, value=item.get(key, ""))
        row += 1

    ws.cell(row=row, column=5, value="合計")
    ws.cell(row=row, column=6, value=f"=SUM(F{start_row + 1}:F{row - 1})")
    ws.cell(row=row, column=5).font = Font(bold=True)
    ws.cell(row=row, column=6).font = Font(bold=True)

    for sheet in [ws, raw]:
        for cells in sheet.iter_rows():
            for cell in cells:
                cell.border = border
                cell.alignment = Alignment(vertical="center", wrap_text=True)

    for index, width in enumerate([18, 38, 16, 12, 14, 14, 14, 18], start=1):
        ws.column_dimensions[get_column_letter(index)].width = width
    ws.freeze_panes = f"A{start_row + 1}"
    ws.auto_filter.ref = f"A{start_row}:F{max(start_row + 1, row - 1)}"

    raw["A1"] = "FAX原文"
    raw["A2"] = payload.get("raw_text", "")
    raw["A1"].font = Font(bold=True)
    raw.column_dimensions["A"].width = 120
    raw.row_dimensions[2].height = 220
    raw["A2"].alignment = Alignment(wrap_text=True, vertical="top")

    wb.save(output_path)
    return output_path


def ensure_ledger_workbook():
    path = ledger_path()
    path.parent.mkdir(exist_ok=True)
    if path.exists():
        wb = load_workbook(path)
        ws = wb["FAX台帳"] if "FAX台帳" in wb.sheetnames else wb.active
        existing_headers = [ws.cell(row=1, column=col).value for col in range(1, ws.max_column + 1)]
        missing_labels = [label for label in LEDGER_COLUMNS.values() if label not in existing_headers]
        if missing_labels:
            start_col = ws.max_column + 1
            header_fill = PatternFill("solid", fgColor="D9EAF7")
            for offset, label in enumerate(missing_labels):
                cell = ws.cell(row=1, column=start_col + offset, value=label)
                cell.fill = header_fill
                cell.font = Font(bold=True)
                ws.column_dimensions[get_column_letter(start_col + offset)].width = 18
            wb.save(path)
        return

    wb = Workbook()
    ws = wb.active
    ws.title = "FAX台帳"
    header_fill = PatternFill("solid", fgColor="D9EAF7")
    thin = Side(style="thin", color="B7C9D6")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    for col, label in enumerate(LEDGER_COLUMNS.values(), start=1):
        cell = ws.cell(row=1, column=col, value=label)
        cell.fill = header_fill
        cell.font = Font(bold=True)
        cell.border = border
        cell.alignment = Alignment(vertical="center", wrap_text=True)
        ws.column_dimensions[get_column_letter(col)].width = 18 if col != 2 else 26
    ws.freeze_panes = "A2"
    wb.save(path)


def append_to_pc_excel(payload, reviewer=None):
    ensure_ledger_workbook()
    path = ledger_path()
    try:
        wb = load_workbook(path)
    except PermissionError as exc:
        raise PermissionError("Excel台帳が開かれています。閉じてからもう一度実行してください。") from exc

    ws = wb["FAX台帳"] if "FAX台帳" in wb.sheetnames else wb.active
    header_to_col = {ws.cell(row=1, column=col).value: col for col in range(1, ws.max_column + 1)}
    header = payload.get("header", {})
    items = payload.get("items") or [{}]
    processed_at = now_text()
    reviewed_at = now_text()
    rows_added = 0

    for item in items:
        row_data = {
            "review_status": "確認済み",
            "reviewer": reviewer or load_config().get("reviewer_default", "社員確認"),
            "reviewed_at": reviewed_at,
            "processed_at": processed_at,
            "source_name": payload.get("source_name", ""),
            **header,
            **item,
        }
        row = ws.max_row + 1
        for key, label in LEDGER_COLUMNS.items():
            col = header_to_col.get(label)
            if not col:
                col = ws.max_column + 1
                ws.cell(row=1, column=col, value=label)
                header_to_col[label] = col
            ws.cell(row=row, column=col, value=row_data.get(key, ""))
        rows_added += 1

    try:
        wb.save(path)
    except PermissionError as exc:
        raise PermissionError("Excel台帳が開かれています。閉じてからもう一度実行してください。") from exc
    write_history("approved_to_excel", {"source": payload.get("source_name", ""), "rows_added": rows_added, "ledger": str(path)})
    return {"path": str(path), "rows_added": rows_added}


def review_id_for(source_name, content_hash):
    base = re.sub(r"[^A-Za-z0-9_-]+", "_", Path(source_name or "manual").stem).strip("_") or "manual"
    return f"{datetime.now().strftime('%Y%m%d_%H%M%S')}_{base}_{content_hash[:8]}"


def save_pending_review(payload, source_name="manual", original_path=""):
    PENDING.mkdir(exist_ok=True)
    raw = payload.get("raw_text", "")
    content_hash = hashlib.sha256((source_name + raw).encode("utf-8", errors="ignore")).hexdigest()

    for existing in PENDING.glob("*.json"):
        try:
            data = json.loads(existing.read_text(encoding="utf-8"))
            if data.get("content_hash") == content_hash:
                return data
        except json.JSONDecodeError:
            continue

    review_id = review_id_for(source_name, content_hash)
    review = {
        "id": review_id,
        "status": "社員確認待ち",
        "created_at": now_text(),
        "source_name": source_name,
        "original_path": original_path,
        "content_hash": content_hash,
        "payload": {**payload, "source_name": source_name},
    }
    (PENDING / f"{review_id}.json").write_text(json.dumps(review, ensure_ascii=False, indent=2), encoding="utf-8")
    write_history("pending_review_created", {"id": review_id, "source": source_name})
    return review


def load_review(review_id):
    path = PENDING / f"{review_id}.json"
    if not path.exists():
        return None
    return json.loads(path.read_text(encoding="utf-8"))


def list_reviews():
    PENDING.mkdir(exist_ok=True)
    reviews = []
    for path in sorted(PENDING.glob("*.json"), key=lambda p: p.stat().st_mtime, reverse=True):
        try:
            review = json.loads(path.read_text(encoding="utf-8"))
            payload = review.get("payload", {})
            header = payload.get("header", {})
            reviews.append(
                {
                    "id": review.get("id"),
                    "status": review.get("status"),
                    "created_at": review.get("created_at"),
                    "source_name": review.get("source_name"),
                    "order_no": header.get("order_no", ""),
                    "project_name": header.get("project_name", ""),
                    "site_name": header.get("site_name", ""),
                    "items_count": len(payload.get("items") or []),
                }
            )
        except json.JSONDecodeError:
            continue
    return reviews


def approve_review(review_id, reviewer=None, payload=None):
    review = load_review(review_id)
    if not review:
        raise FileNotFoundError("確認待ちデータが見つかりません。")
    final_payload = payload or review.get("payload", {})
    result = append_to_pc_excel(final_payload, reviewer=reviewer)
    (PENDING / f"{review_id}.json").unlink(missing_ok=True)
    write_history("review_approved", {"id": review_id, "reviewer": reviewer or ""})
    return result


def reject_review(review_id, reason=""):
    review = load_review(review_id)
    if not review:
        raise FileNotFoundError("確認待ちデータが見つかりません。")
    REJECTED.mkdir(exist_ok=True)
    review["status"] = "差し戻し"
    review["rejected_at"] = now_text()
    review["reject_reason"] = reason
    (REJECTED / f"{review_id}.json").write_text(json.dumps(review, ensure_ascii=False, indent=2), encoding="utf-8")
    (PENDING / f"{review_id}.json").unlink(missing_ok=True)
    write_history("review_rejected", {"id": review_id, "reason": reason})
    return {"id": review_id, "status": "差し戻し"}


def process_incoming_once():
    folder = incoming_dir()
    folder.mkdir(exist_ok=True)
    PROCESSED.mkdir(exist_ok=True)
    results = []
    for path in sorted(p for p in folder.iterdir() if p.is_file()):
        file_bytes = path.read_bytes()
        text, warning = read_file_text(path.name, file_bytes)
        if not text:
            results.append({"source": path.name, "status": "needs_ocr", "warning": warning})
            continue
        data = extract_fax(text)
        data["source_name"] = path.stem
        review = save_pending_review(data, source_name=path.name, original_path=str(path))
        target = PROCESSED / path.name
        if target.exists():
            target = PROCESSED / f"{path.stem}_{datetime.now().strftime('%H%M%S')}{path.suffix}"
        path.replace(target)
        results.append({"source": path.name, "status": "pending_review", "review_id": review["id"], "warning": warning})
    return results


class Handler(BaseHTTPRequestHandler):
    def send_json(self, data, status=200):
        body = json.dumps(data, ensure_ascii=False).encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def do_GET(self):
        parsed = urlparse(self.path)
        if parsed.path == "/health":
            self.send_json({"status": "ok", "version": APP_VERSION, "config": load_config()})
            return
        if parsed.path == "/api/config":
            config = load_config()
            self.send_json({**config, "pending_count": len(list_reviews())})
            return
        if parsed.path == "/api/reviews":
            self.send_json({"reviews": list_reviews()})
            return
        if parsed.path.startswith("/api/reviews/"):
            review = load_review(Path(parsed.path).name)
            if not review:
                self.send_json({"error": "確認待ちデータが見つかりません。"}, 404)
                return
            self.send_json(review)
            return
        if parsed.path.startswith("/download/"):
            self.send_download(Path(parsed.path).name)
            return

        path = "/index.html" if parsed.path == "/" else parsed.path
        file_path = (PUBLIC / path.lstrip("/")).resolve()
        if not str(file_path).startswith(str(PUBLIC.resolve())) or not file_path.exists():
            self.send_error(404)
            return
        mime = "text/html; charset=utf-8"
        if file_path.suffix == ".css":
            mime = "text/css; charset=utf-8"
        elif file_path.suffix == ".js":
            mime = "application/javascript; charset=utf-8"
        body = file_path.read_bytes()
        self.send_response(200)
        self.send_header("Content-Type", mime)
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def multipart_files(self):
        ctype, _ = cgi.parse_header(self.headers.get("Content-Type", ""))
        if ctype != "multipart/form-data":
            return []
        form = cgi.FieldStorage(fp=self.rfile, headers=self.headers, environ={"REQUEST_METHOD": "POST"})
        if "file" not in form:
            return []
        item = form["file"]
        return item if isinstance(item, list) else [item]

    def read_json_body(self):
        length = int(self.headers.get("Content-Length", "0"))
        return json.loads(self.rfile.read(length).decode("utf-8") or "{}")

    def do_POST(self):
        if self.path in {"/api/read-file", "/api/read-pdf"}:
            started = time.monotonic()
            config = load_config()
            request_timeout = float(config.get("request_timeout_seconds", 9))
            max_upload_files = int(config.get("max_upload_files", 1))
            file_items = self.multipart_files()
            if not file_items:
                self.send_json({"error": "ファイルを読み取れませんでした。"}, 400)
                return
            parts = []
            warnings = []
            if len(file_items) > max_upload_files:
                warnings.append(f"10秒以内にするため先頭{max_upload_files}ファイルだけ読みました。")
            for file_item in file_items[:max_upload_files]:
                if time.monotonic() - started > request_timeout:
                    warnings.append("10秒以内にするため、残りの読み取りを打ち切りました。")
                    break
                if file_item is None or not getattr(file_item, "file", None):
                    continue
                filename = file_item.filename or "upload"
                text, warning = read_file_text(filename, file_item.file.read())
                if text:
                    parts.append(f"--- {filename} ---\n{text}")
                if warning:
                    warnings.append(f"{filename}: {warning}")
                if time.monotonic() - started > request_timeout:
                    warnings.append("10秒以内にするため、ここで読み取りを終了しました。")
                    break
            self.send_json(
                {
                    "text": "\n\n".join(parts).strip(),
                    "warning": " / ".join(warnings),
                    "elapsed_seconds": round(time.monotonic() - started, 2),
                    "version": APP_VERSION,
                    "mode": "20秒以内・精度優先",
                }
            )
            return

        if self.path == "/api/auto-run":
            self.send_json({"folder": str(incoming_dir()), "results": process_incoming_once()})
            return

        try:
            data = self.read_json_body()
        except json.JSONDecodeError:
            self.send_json({"error": "JSONを読み取れませんでした。"}, 400)
            return

        if self.path == "/api/config":
            config = load_config()
            for key in [
                "incoming_fax_dir",
                "ledger_path",
                "reviewer_default",
                "ocr_min_chars",
                "ocr_psm",
                "ocr_timeout_seconds",
                "ocr_preprocess",
                "ocr_preprocess_mode",
                "ocr_processed_format",
                "ocr_jpeg_quality",
                "ocr_pdf_dpi",
                "ocr_pdf_image_format",
                "ocr_target_long_side",
                "ocr_min_long_side",
                "ocr_max_pages",
                "ocr_retry_when_empty",
                "request_timeout_seconds",
                "max_upload_files",
                "pdf_text_max_pages",
                "xlsx_max_rows_per_sheet",
                "docx_max_paragraphs",
                "docx_max_table_rows",
                "html_max_chars",
                "email_max_parts",
                "zip_max_files",
                "zip_max_file_bytes",
                "text_max_chars",
            ]:
                if key in data:
                    config[key] = data[key]
            save_config(config)
            self.send_json(config)
            return

        if self.path == "/api/extract":
            self.send_json(extract_fax(data.get("text", "")))
            return

        if self.path == "/api/export":
            output_path = create_workbook(data)
            self.send_json({"file": output_path.name, "download_url": f"/download/{output_path.name}"})
            return

        if self.path == "/api/save-review":
            review = save_pending_review(data, source_name=data.get("source_name", "manual"))
            self.send_json(review)
            return

        if self.path == "/api/approve-review":
            try:
                result = approve_review(data.get("id", ""), reviewer=data.get("reviewer"), payload=data.get("payload"))
            except (FileNotFoundError, PermissionError) as exc:
                self.send_json({"error": str(exc)}, 409)
                return
            self.send_json(result)
            return

        if self.path == "/api/reject-review":
            try:
                result = reject_review(data.get("id", ""), reason=data.get("reason", ""))
            except FileNotFoundError as exc:
                self.send_json({"error": str(exc)}, 404)
                return
            self.send_json(result)
            return

        if self.path == "/api/append-ledger":
            try:
                result = append_to_pc_excel(data, reviewer=data.get("reviewer"))
            except PermissionError as exc:
                self.send_json({"error": str(exc)}, 409)
                return
            self.send_json(result)
            return

        self.send_error(404)

    def send_download(self, filename):
        file_path = (OUTPUTS / filename).resolve()
        if not str(file_path).startswith(str(OUTPUTS.resolve())) or not file_path.exists():
            self.send_error(404)
            return
        body = file_path.read_bytes()
        self.send_response(200)
        self.send_header("Content-Type", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        self.send_header("Content-Disposition", f'attachment; filename="{filename}"')
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def log_message(self, format, *args):
        return


def main():
    for folder in [incoming_dir(), OUTPUTS, PENDING, PROCESSED, REJECTED, PC_EXCEL_DIR]:
        folder.mkdir(exist_ok=True)
    port = int(os.environ.get("PORT") or (sys.argv[1] if len(sys.argv) > 1 else 8765))
    host = os.environ.get("HOST", "0.0.0.0")
    server = ThreadingHTTPServer((host, port), Handler)
    print(f"http://{host}:{port}")
    server.serve_forever()


if __name__ == "__main__":
    main()
